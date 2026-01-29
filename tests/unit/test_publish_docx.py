"""Unit tests for DOCX publishing."""

import io
import tempfile
from pathlib import Path

from docx import Document

from jamb.core.models import Item, TraceabilityGraph
from jamb.publish.formats.docx import generate_template, render_docx


class TestRenderDocx:
    """Tests for the render_docx function."""

    def test_render_docx_returns_bytes(self):
        """Test that render_docx returns bytes."""
        items = [
            Item(uid="SRS001", text="Test requirement", document_prefix="SRS"),
        ]
        result = render_docx(items, "SRS")

        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_render_docx_valid_document(self):
        """Test that render_docx produces a valid DOCX file."""
        items = [
            Item(uid="SRS001", text="Test requirement", document_prefix="SRS"),
        ]
        result = render_docx(items, "SRS")

        # Should be able to open as a valid DOCX
        doc = Document(io.BytesIO(result))
        assert doc is not None

    def test_render_docx_contains_title(self):
        """Test that render_docx includes document title."""
        items = [
            Item(uid="SRS001", text="Test requirement", document_prefix="SRS"),
        ]
        result = render_docx(items, "SRS")

        doc = Document(io.BytesIO(result))
        # First paragraph should contain the document title
        paragraphs = [p.text for p in doc.paragraphs]
        assert any("SRS" in p for p in paragraphs)

    def test_render_docx_contains_item_uid(self):
        """Test that render_docx includes item UIDs."""
        items = [
            Item(uid="SRS001", text="Test requirement 1", document_prefix="SRS"),
            Item(uid="SRS002", text="Test requirement 2", document_prefix="SRS"),
        ]
        result = render_docx(items, "SRS")

        doc = Document(io.BytesIO(result))
        paragraphs = [p.text for p in doc.paragraphs]
        full_text = " ".join(paragraphs)

        assert "SRS001" in full_text
        assert "SRS002" in full_text

    def test_render_docx_contains_item_text(self):
        """Test that render_docx includes item text."""
        items = [
            Item(
                uid="SRS001",
                text="The system shall authenticate users.",
                document_prefix="SRS",
            ),
        ]
        result = render_docx(items, "SRS")

        doc = Document(io.BytesIO(result))
        paragraphs = [p.text for p in doc.paragraphs]
        full_text = " ".join(paragraphs)

        assert "authenticate users" in full_text

    def test_render_docx_contains_headers(self):
        """Test that render_docx includes item headers."""
        items = [
            Item(
                uid="SRS001",
                text="Some text",
                header="Authentication Requirement",
                document_prefix="SRS",
            ),
        ]
        result = render_docx(items, "SRS")

        doc = Document(io.BytesIO(result))
        paragraphs = [p.text for p in doc.paragraphs]
        full_text = " ".join(paragraphs)

        assert "Authentication Requirement" in full_text

    def test_render_docx_contains_links(self):
        """Test that render_docx includes item links."""
        items = [
            Item(
                uid="SRS001",
                text="Test requirement",
                document_prefix="SRS",
                links=["UN001", "UN002"],
            ),
        ]
        result = render_docx(items, "SRS")

        doc = Document(io.BytesIO(result))
        paragraphs = [p.text for p in doc.paragraphs]
        full_text = " ".join(paragraphs)

        assert "Links:" in full_text
        assert "UN001" in full_text
        assert "UN002" in full_text

    def test_render_docx_no_links_when_disabled(self):
        """Test that render_docx omits links when include_links is False."""
        items = [
            Item(
                uid="SRS001",
                text="Test requirement",
                document_prefix="SRS",
                links=["UN001"],
            ),
        ]
        result = render_docx(items, "SRS", include_links=False)

        doc = Document(io.BytesIO(result))
        paragraphs = [p.text for p in doc.paragraphs]
        full_text = " ".join(paragraphs)

        assert "Links:" not in full_text

    def test_render_docx_empty_items_list(self):
        """Test that render_docx handles empty items list."""
        result = render_docx([], "SRS")

        assert isinstance(result, bytes)
        doc = Document(io.BytesIO(result))
        # Should still have title and summary
        paragraphs = [p.text for p in doc.paragraphs]
        full_text = " ".join(paragraphs)
        assert "Total items: 0" in full_text

    def test_render_docx_sorts_by_uid(self):
        """Test that render_docx sorts items by UID."""
        items = [
            Item(uid="SRS002", text="Second", document_prefix="SRS"),
            Item(uid="SRS001", text="First", document_prefix="SRS"),
            Item(uid="SRS003", text="Third", document_prefix="SRS"),
        ]
        result = render_docx(items, "SRS")

        doc = Document(io.BytesIO(result))
        paragraphs = [p.text for p in doc.paragraphs]
        full_text = " ".join(paragraphs)

        # Find positions of each item
        pos_001 = full_text.find("SRS001")
        pos_002 = full_text.find("SRS002")
        pos_003 = full_text.find("SRS003")

        # Should be sorted by UID
        assert pos_001 < pos_002 < pos_003

    def test_render_docx_item_without_text(self):
        """Test that render_docx handles items without text."""
        items = [
            Item(uid="SRS001", text="", document_prefix="SRS", header="Header Only"),
        ]
        result = render_docx(items, "SRS")

        # Should not raise an error
        assert isinstance(result, bytes)
        doc = Document(io.BytesIO(result))
        assert doc is not None

    def test_render_docx_multiple_documents(self):
        """Test that render_docx handles items from multiple documents."""
        items = [
            Item(uid="UN001", text="Customer need", document_prefix="UN"),
            Item(uid="SRS001", text="Software req", document_prefix="SRS"),
            Item(uid="SYS001", text="System req", document_prefix="SYS"),
        ]
        result = render_docx(items, "Requirements Document")

        doc = Document(io.BytesIO(result))
        paragraphs = [p.text for p in doc.paragraphs]
        full_text = " ".join(paragraphs)

        # Should contain all items
        assert "UN001" in full_text
        assert "SRS001" in full_text
        assert "SYS001" in full_text

    def test_render_docx_links_as_hyperlinks(self):
        """Test that links to items in the document become hyperlinks."""
        items = [
            Item(uid="UN001", text="Customer need", document_prefix="UN"),
            Item(
                uid="SRS001",
                text="Software req",
                document_prefix="SRS",
                links=["UN001"],
            ),
        ]
        result = render_docx(items, "Requirements Document")

        # Should produce valid docx with hyperlinks
        doc = Document(io.BytesIO(result))
        assert doc is not None

        # Check that the document contains the link text
        paragraphs = [p.text for p in doc.paragraphs]
        full_text = " ".join(paragraphs)
        assert "Links:" in full_text

    def test_render_docx_child_links_with_graph(self):
        """Test that child links appear when graph is provided."""
        graph = TraceabilityGraph()
        un_item = Item(uid="UN001", text="Customer need", document_prefix="UN")
        srs_item = Item(
            uid="SRS001",
            text="Software req",
            document_prefix="SRS",
            links=["UN001"],
        )
        graph.add_item(un_item)
        graph.add_item(srs_item)

        items = [un_item, srs_item]
        result = render_docx(items, "Requirements Document", graph=graph)

        doc = Document(io.BytesIO(result))
        paragraphs = [p.text for p in doc.paragraphs]
        full_text = " ".join(paragraphs)

        # UN001 should have "Linked from:" showing SRS001
        assert "Linked from:" in full_text

    def test_render_docx_heading_item_uses_level1(self):
        """Test that heading items use level 1 heading."""
        items = [
            Item(
                uid="SRS001",
                text="",
                header="Safety Requirements",
                document_prefix="SRS",
                type="heading",
            ),
        ]
        result = render_docx(items, "SRS")

        doc = Document(io.BytesIO(result))
        paragraphs = [p.text for p in doc.paragraphs]
        full_text = " ".join(paragraphs)

        # Should contain the header text
        assert "Safety Requirements" in full_text
        # Should NOT have UID prefix in heading
        assert "SRS001: Safety Requirements" not in full_text

    def test_render_docx_info_item_italic(self):
        """Test that info items have italic paragraph text."""
        items = [
            Item(
                uid="SRS002",
                text="This is informational text.",
                document_prefix="SRS",
                type="info",
            ),
        ]
        result = render_docx(items, "SRS")

        doc = Document(io.BytesIO(result))
        # Find the paragraph with the info text
        info_para = None
        for p in doc.paragraphs:
            if "This is informational text." in p.text:
                info_para = p
                break

        assert info_para is not None
        # Check that the run is italic
        assert info_para.runs[0].italic is True

    def test_render_docx_child_links_no_graph(self):
        """Child links not shown when graph is None."""
        items = [
            Item(uid="UN001", text="Need", document_prefix="UN"),
            Item(uid="SRS001", text="Req", document_prefix="SRS", links=["UN001"]),
        ]
        result = render_docx(items, "Test", graph=None)
        doc = Document(io.BytesIO(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Linked from:" not in full_text

    def test_render_docx_external_link_plain_text(self):
        """Links to items not in the document appear as plain text."""
        items = [
            Item(
                uid="SRS001",
                text="Req",
                document_prefix="SRS",
                links=["EXTERNAL001"],
            ),
        ]
        result = render_docx(items, "SRS")
        doc = Document(io.BytesIO(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "EXTERNAL001" in full_text

    def test_render_docx_external_child_links_not_shown(self):
        """Child links to items not in rendered set are not shown."""
        graph = TraceabilityGraph()
        un_item = Item(uid="UN001", text="Need", document_prefix="UN")
        srs_item = Item(uid="SRS001", text="Req", document_prefix="SRS", links=["UN001"])
        graph.add_item(un_item)
        graph.add_item(srs_item)

        # Only render UN001
        items = [un_item]
        result = render_docx(items, "Test", graph=graph)
        doc = Document(io.BytesIO(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Linked from:" not in full_text

    def test_render_docx_heading_no_header_falls_back_to_uid(self):
        """6d: Heading item with no header falls back to UID."""
        items = [
            Item(
                uid="SRS001",
                text="",
                document_prefix="SRS",
                type="heading",
                header="",
            ),
        ]
        result = render_docx(items, "SRS")
        doc = Document(io.BytesIO(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "SRS001" in full_text

    def test_render_docx_unicode_text(self):
        """6e: Unicode characters in item text produce no encoding errors."""
        items = [
            Item(
                uid="SRS001",
                text="Ünïcödé — «text» ñ 日本語",
                document_prefix="SRS",
            ),
        ]
        result = render_docx(items, "SRS")
        doc = Document(io.BytesIO(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Ünïcödé" in full_text
        assert "日本語" in full_text

    def test_links_not_in_all_uids_plain_text(self):
        items = [
            Item(uid="SRS001", text="Req", document_prefix="SRS", links=["PHANTOM001"]),
        ]
        result = render_docx(items, "Test")
        doc = Document(io.BytesIO(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "PHANTOM001" in full_text
        assert "Links:" in full_text
        # Find links paragraph and verify no hyperlink element for phantom
        for p in doc.paragraphs:
            if "Links:" in p.text:
                xml_str = p._p.xml
                assert 'w:anchor="PHANTOM001"' not in xml_str
                break

    def test_all_unknown_document_order(self):
        items = [
            Item(uid="ZZZ001", text="Zeta", document_prefix="ZZZ"),
            Item(uid="AAA001", text="Alpha", document_prefix="AAA"),
        ]
        result = render_docx(items, "Test", document_order=["UN"])
        doc = Document(io.BytesIO(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert full_text.index("AAA001") < full_text.index("ZZZ001")

    def test_empty_links_list_no_section(self):
        items = [
            Item(uid="SRS001", text="Req", document_prefix="SRS", links=[]),
        ]
        result = render_docx(items, "Test")
        doc = Document(io.BytesIO(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Links:" not in full_text


class TestRenderDocxEdgeCases:
    """Additional edge case tests for render_docx."""

    def test_include_links_false_with_graph_suppresses_all_links(self):
        """include_links=False with a graph suppresses
        both parent links and child links."""
        graph = TraceabilityGraph()
        un_item = Item(uid="UN001", text="Customer need", document_prefix="UN")
        srs_item = Item(
            uid="SRS001",
            text="Software req",
            document_prefix="SRS",
            links=["UN001"],
        )
        graph.add_item(un_item)
        graph.add_item(srs_item)

        items = [un_item, srs_item]
        result = render_docx(
            items,
            "Requirements",
            include_links=False,
            graph=graph,
        )

        doc = Document(io.BytesIO(result))
        full_text = " ".join(p.text for p in doc.paragraphs)

        # Neither parent links nor reverse child links should appear
        assert "Links:" not in full_text
        assert "Linked from:" not in full_text

    def test_document_order_sorts_items_by_hierarchy(self):
        """Items are sorted according to document_order parameter."""
        items = [
            Item(uid="SRS001", text="Software req", document_prefix="SRS"),
            Item(uid="UN001", text="User need", document_prefix="UN"),
            Item(uid="SYS001", text="System req", document_prefix="SYS"),
        ]
        result = render_docx(
            items,
            "Full Hierarchy",
            document_order=["UN", "SYS", "SRS"],
        )

        doc = Document(io.BytesIO(result))
        full_text = " ".join(p.text for p in doc.paragraphs)

        # UN should come before SYS, SYS before SRS
        un_pos = full_text.find("UN001")
        sys_pos = full_text.find("SYS001")
        srs_pos = full_text.find("SRS001")
        assert un_pos < sys_pos < srs_pos


class TestRenderDocxTemplateSupport:
    """Tests for DOCX template support."""

    def test_render_docx_with_template(self):
        """Test that render_docx works with a template file."""
        # First generate a template
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            template_path = f.name

        try:
            generate_template(template_path)

            # Now use it to render items
            items = [
                Item(uid="SRS001", text="Test requirement", document_prefix="SRS"),
            ]
            result = render_docx(items, "SRS", template_path=template_path)

            # Should produce valid docx
            doc = Document(io.BytesIO(result))
            full_text = " ".join(p.text for p in doc.paragraphs)
            assert "SRS001" in full_text
            assert "Test requirement" in full_text
        finally:
            Path(template_path).unlink(missing_ok=True)

    def test_render_docx_template_clears_content(self):
        """Test that template content is cleared before adding new items."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            template_path = f.name

        try:
            generate_template(template_path)

            items = [
                Item(uid="NEW001", text="New item", document_prefix="NEW"),
            ]
            result = render_docx(items, "New Doc", template_path=template_path)

            doc = Document(io.BytesIO(result))
            full_text = " ".join(p.text for p in doc.paragraphs)

            # Should NOT contain template example content
            assert "How to use this template" not in full_text
            # Should contain new items
            assert "NEW001" in full_text
        finally:
            Path(template_path).unlink(missing_ok=True)


class TestGenerateTemplate:
    """Tests for the template generator function."""

    def test_generate_template_creates_file(self):
        """Test that generate_template creates a valid DOCX file."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name

        try:
            generate_template(output_path)

            # File should exist and be valid docx
            assert Path(output_path).exists()
            doc = Document(output_path)
            assert doc is not None
        finally:
            Path(output_path).unlink(missing_ok=True)

    def test_generate_template_contains_sample_content(self):
        """Test that generated template has sample content."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name

        try:
            generate_template(output_path)
            doc = Document(output_path)
            full_text = " ".join(p.text for p in doc.paragraphs)

            # Should have explanatory content
            assert "Jamb Document Template" in full_text
            assert "How to use this template" in full_text
            assert "--template" in full_text
        finally:
            Path(output_path).unlink(missing_ok=True)

    def test_generate_template_has_style_examples(self):
        """Test that generated template shows different styles."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name

        try:
            generate_template(output_path)
            doc = Document(output_path)
            full_text = " ".join(p.text for p in doc.paragraphs)

            # Should have examples of different styles
            assert "Heading 1" in full_text
            assert "Heading 2" in full_text
            assert "Normal" in full_text
        finally:
            Path(output_path).unlink(missing_ok=True)


class TestRenderDocxPageNumbers:
    """Tests for page number functionality."""

    def test_render_docx_has_footer(self):
        """Test that rendered document has footer sections."""
        items = [
            Item(uid="SRS001", text="Test requirement", document_prefix="SRS"),
        ]
        result = render_docx(items, "SRS")

        doc = Document(io.BytesIO(result))
        # Document should have at least one section with footer
        assert len(doc.sections) > 0
        section = doc.sections[0]
        # Footer should exist (even if empty when not rendered)
        assert section.footer is not None
