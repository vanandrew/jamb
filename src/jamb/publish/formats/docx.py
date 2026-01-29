"""Word (DOCX) document output for publishing."""

import io
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

from jamb.core.models import Item, TraceabilityGraph


def _add_bookmark(paragraph: Paragraph, bookmark_name: str, bookmark_id: int) -> None:
    """Add a bookmark to a paragraph for internal linking.

    Args:
        paragraph: The docx paragraph object to add the bookmark to.
        bookmark_name: The bookmark identifier used for internal
            cross-references.
        bookmark_id: Unique numeric ID for the bookmark element.
    """
    # Create bookmark start
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(bookmark_id))
    bookmark_start.set(qn("w:name"), bookmark_name)

    # Create bookmark end
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(bookmark_id))

    # Insert into paragraph
    paragraph._p.insert(0, bookmark_start)
    paragraph._p.append(bookmark_end)


def _add_hyperlink(paragraph: Paragraph, anchor: str, text: str) -> None:
    """Add an internal hyperlink to a paragraph.

    Args:
        paragraph: The docx paragraph object to add the hyperlink to.
        anchor: The bookmark name to link to.
        text: The display text for the hyperlink.
    """
    # Create hyperlink element
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)

    # Create run with text
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")

    # Add hyperlink styling (blue, underlined)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    run_props.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(underline)

    run.append(run_props)

    # Add text
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_page_numbers(doc: Any) -> None:
    """Add page numbers to the document footer.

    Args:
        doc: The Document object to add page numbers to.
    """
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add "Page " text
        run = para.add_run("Page ")
        run.font.size = Pt(9)

        # Add PAGE field
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.text = "PAGE"

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run2 = para.add_run()
        run2._r.append(fld_char_begin)
        run2._r.append(instr_text)
        run2._r.append(fld_char_end)
        run2.font.size = Pt(9)

        # Add " of " text
        run3 = para.add_run(" of ")
        run3.font.size = Pt(9)

        # Add NUMPAGES field
        fld_char_begin2 = OxmlElement("w:fldChar")
        fld_char_begin2.set(qn("w:fldCharType"), "begin")

        instr_text2 = OxmlElement("w:instrText")
        instr_text2.text = "NUMPAGES"

        fld_char_end2 = OxmlElement("w:fldChar")
        fld_char_end2.set(qn("w:fldCharType"), "end")

        run4 = para.add_run()
        run4._r.append(fld_char_begin2)
        run4._r.append(instr_text2)
        run4._r.append(fld_char_end2)
        run4.font.size = Pt(9)


def _setup_styles(doc: Any) -> None:
    """Set up document styles for better formatting.

    Args:
        doc: The Document object to configure styles for.
    """
    styles = doc.styles

    # Configure Normal style
    if "Normal" in styles:
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(8)
        normal.paragraph_format.line_spacing = 1.15

    # Configure Heading 1
    if "Heading 1" in styles:
        h1 = styles["Heading 1"]
        h1.font.name = "Calibri"
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(6)

    # Configure Heading 2
    if "Heading 2" in styles:
        h2 = styles["Heading 2"]
        h2.font.name = "Calibri"
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0x2C, 0x52, 0x82)
        h2.paragraph_format.space_before = Pt(14)
        h2.paragraph_format.space_after = Pt(4)

    # Create or update Item Type style for type indicators
    if "Intense Emphasis" not in styles:
        item_type_style = styles.add_style("Intense Emphasis", WD_STYLE_TYPE.CHARACTER)
    else:
        item_type_style = styles["Intense Emphasis"]
    item_type_style.font.bold = True
    item_type_style.font.size = Pt(9)
    item_type_style.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)


def render_docx(
    items: list[Item],
    title: str,
    include_links: bool = True,
    document_order: list[str] | None = None,
    graph: TraceabilityGraph | None = None,
    template_path: str | None = None,
) -> bytes:
    """Render items as a Word document.

    Args:
        items: List of Item objects to include in the document.
        title: The document title.
        include_links: Whether to include links section for each item.
        document_order: Optional list of document prefixes in hierarchy order
                       (root first, then children). If None, sorts alphabetically.
        graph: Optional traceability graph for child link lookup.
        template_path: Optional path to a DOCX template file to use as base.

    Returns:
        Bytes representing the DOCX file.
    """
    if template_path:
        doc = Document(template_path)
        # Clear existing content from template (keep styles)
        for element in doc.element.body[:]:
            doc.element.body.remove(element)
    else:
        doc = Document()
        _setup_styles(doc)

    # Add page numbers to footer
    _add_page_numbers(doc)

    # Build set of all item UIDs for link validation
    all_uids = {item.uid for item in items}

    # Add title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add summary paragraph
    doc.add_paragraph(f"Total items: {len(items)}")
    doc.add_paragraph()

    # Build document order index for sorting
    if document_order:
        doc_order_index = {prefix: i for i, prefix in enumerate(document_order)}
    else:
        doc_order_index = {}

    fallback_order = len(document_order) if document_order else 0

    def get_doc_order(item: Item) -> int:
        return doc_order_index.get(item.document_prefix, fallback_order)

    # Sort items by document hierarchy, then UID
    sorted_items = sorted(items, key=lambda x: (get_doc_order(x), x.document_prefix, x.uid))

    # Track current document for section headers
    current_doc = None

    for idx, item in enumerate(sorted_items):
        # Add document section header if document changed
        if item.document_prefix != current_doc:
            current_doc = item.document_prefix
            doc.add_heading(f"{current_doc}", level=1)

        # Create heading with UID and optional header
        item_type = item.type

        if item_type == "heading":
            heading_display = item.header if item.header else item.uid
            heading = doc.add_heading(heading_display, level=1)
        else:
            if item.header:
                heading_text = f"{item.uid}: {item.header}"
            else:
                heading_text = item.uid
            heading = doc.add_heading(heading_text, level=2)

        # Add bookmark for this item so links can reference it
        _add_bookmark(heading, item.uid, idx + 100)

        # Add item text
        if item.text:
            if item_type == "info":
                para = doc.add_paragraph()
                run = para.add_run(item.text)
                run.italic = True
                run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
            else:
                doc.add_paragraph(item.text)

        # Add links section with hyperlinks
        if include_links and item.links:
            links_para = doc.add_paragraph()
            links_run = links_para.add_run("Links: ")
            links_run.bold = True

            for i, link_uid in enumerate(item.links):
                if i > 0:
                    links_para.add_run(", ")

                # Add as hyperlink if target exists, otherwise plain text
                if link_uid in all_uids:
                    _add_hyperlink(links_para, link_uid, link_uid)
                else:
                    links_para.add_run(link_uid)

        # Add child links section (reverse links)
        if include_links and graph is not None:
            children = graph.item_children.get(item.uid, [])
            visible_children = [c for c in children if c in all_uids]
            if visible_children:
                child_para = doc.add_paragraph()
                child_run = child_para.add_run("Linked from: ")
                child_run.bold = True

                for i, child_uid in enumerate(visible_children):
                    if i > 0:
                        child_para.add_run(", ")
                    _add_hyperlink(child_para, child_uid, child_uid)

        # Add spacing between items
        doc.add_paragraph()

    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def generate_template(output_path: str) -> None:
    """Generate a sample DOCX template file with jamb styles.

    Creates a template document that users can customize and then use
    with the --template option when publishing documents.

    Args:
        output_path: Path where the template file will be written.
    """
    doc = Document()
    _setup_styles(doc)

    # Add title explaining the template
    title = doc.add_heading("Jamb Document Template", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.add_run("How to use this template:").bold = True

    doc.add_paragraph(
        "This template defines styles that jamb uses when publishing documents. "
        "You can customize fonts, colors, and spacing by modifying the styles "
        "in this document. Then use the --template option when publishing:"
    )

    code_para = doc.add_paragraph()
    code_run = code_para.add_run("    jamb publish SRS output.docx --template this-file.docx")
    code_run.italic = True
    code_run.font.size = Pt(10)

    doc.add_paragraph()

    # Section showing Heading 1 style (used for document sections and heading items)
    doc.add_heading("Document Section (Heading 1 Style)", level=1)
    doc.add_paragraph(
        "This style is used for document section headers (e.g., 'SRS', 'UN') "
        "and for items with type='heading'. Customize this style to change "
        "the appearance of major sections in your published document."
    )

    doc.add_paragraph()

    # Section showing Heading 2 style (used for requirement items)
    doc.add_heading("SRS001: Requirement Item (Heading 2 Style)", level=2)
    doc.add_paragraph(
        "This style is used for requirement item headings. Each requirement "
        "shows its UID and optional header text using this style."
    )

    doc.add_paragraph()

    # Show normal paragraph style
    doc.add_heading("SRS002: Another Requirement", level=2)
    doc.add_paragraph(
        "This is the Normal paragraph style, used for requirement body text. "
        "Customize font, size, and spacing to match your organization's "
        "document standards."
    )

    doc.add_paragraph()

    # Show info item style (italic)
    doc.add_heading("SRS003: Info Item Example", level=2)
    info_para = doc.add_paragraph()
    info_run = info_para.add_run(
        "Info items are displayed in italic with muted color. This style is "
        "useful for explanatory notes that aren't formal requirements."
    )
    info_run.italic = True
    info_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    doc.add_paragraph()

    # Show links style
    doc.add_heading("SRS004: Item With Links", level=2)
    doc.add_paragraph("The system shall do something important.")
    links_para = doc.add_paragraph()
    links_run = links_para.add_run("Links: ")
    links_run.bold = True
    links_para.add_run("SYS001, SYS002")

    child_para = doc.add_paragraph()
    child_run = child_para.add_run("Linked from: ")
    child_run.bold = True
    child_para.add_run("UT001, UT002")

    doc.add_paragraph()

    # Tips section
    doc.add_heading("Tips for Customization", level=1)

    tips = [
        "Open this template in Microsoft Word or compatible editor",
        "Modify styles via the Styles pane (don't just format text directly)",
        "Heading 1: Document sections and heading-type items",
        "Heading 2: Requirement item headings",
        "Normal: Body text for requirements",
        "Save your customized template and use with --template",
    ]

    for tip in tips:
        doc.add_paragraph(tip, style="List Bullet")

    # Add page numbers to show they work
    _add_page_numbers(doc)

    # Save the template
    doc.save(output_path)
